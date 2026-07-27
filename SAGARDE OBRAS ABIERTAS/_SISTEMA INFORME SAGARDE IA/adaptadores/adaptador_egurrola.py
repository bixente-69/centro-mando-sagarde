# -*- coding: utf-8 -*-
"""
ADAPTADOR - 2025 GETXO 12V EGURROLA  (capa 1 - lectura, especifico de obra)
------------------------------------------------------------------------------
Sabe leer el formato concreto de la hoja de revisiones de esta obra: MAS
SIMPLE que Mungia y que Obispo Orueta.

  - Los 3 ficheros de revision estan en la RAIZ de la carpeta de la obra
    (no hay subcarpeta "REVISIONES"): "REVISION EGURROLA .docx" (con un
    espacio antes del punto), "REVISION EGURROLA 2.docx" y
    "REVISION EGURROLA 3.docx". Se localizan por prefijo de nombre
    ("revision egurrola"), no por una lista fija, para no romperse si se
    anade una 4a revision en el futuro.

  - UNA sola tabla por documento, que en realidad cubre las 4 plantas de
    la obra (Baja/1/2/3) en DOS "bloques" verticales de 2 columnas dobles
    cada uno (izquierda/derecha = 2 plantas en paralelo, igual que en
    Mungia pero sin tanta anidacion): bloque 1 = plantas 0 y 1 (filas
    1-31), bloque 2 = plantas 2 y 3 (filas 35-65). Cada bloque empieza con
    una fila de cabecera "TRABAJO | PLANTA | A | B | C | TRABAJO | PLANTA
    | A | B | C" que da los codigos reales de vivienda (aqui siempre A/B/C,
    pero se LEEN de la cabecera en vez de darse por supuestos, por si
    cambiaran). Confirmado a mano sobre los 3 ficheros: 4 plantas x 3
    viviendas/planta = 12 viviendas, que coincide con "12 VIVIENDAS EN
    CALLE ARTIBAI" (portada de "2025 12v egurrola getxo.pdf") y con las
    4 plantas reales del proyecto (PLANTA BAJA/PRIMERA/SEGUNDA/TERCERA,
    vistas en ese mismo PDF), asi que floor '0' se relabela a 'PB' y el
    resto se deja tal cual ('1','2','3').

  - Un UNICO edificio/portal (no hay ZR1.1/ZR1.2 ni varios portales
    detectados en ningun documento de la obra). Se usa BUILDING = "Artibai"
    como codigo interno. OJO: el numero de portal exacto NO esta
    confirmado con certeza: "TIERRAS EGURROLA.pdf" cita como LOCAL
    "C. Artibai, 3C, 48991 Getxo, Vizcaya", pero "3C" es ambiguo -- podria
    ser "portal 3, puerta C" o, igual de plausible, "planta 3 / vivienda C"
    (la misma nomenclatura planta+letra que usa la propia tabla de
    revisiones). Se deja como [pendiente de confirmar] en la Ficha de Obra
    en vez de asumir que el portal es el nº 3.

  - IMPORTANTE - columnas 10 vs 12: "REVISION EGURROLA .docx" (la de sin
    numero) tiene 12 columnas (2 columnas vacias de separacion visual entre
    el bloque izquierdo y el derecho); "REVISION EGURROLA 2.docx" y
    "REVISION EGURROLA 3.docx" tienen 10 (sin esas 2 columnas vacias). El
    parseo NO asume un numero fijo de columnas: el bloque derecho
    (TRABAJO/PLANTA/A/B/C) siempre son las ULTIMAS 5 columnas de la fila,
    sea la tabla de 10 o de 12, asi que se localiza como
    `fila[len(fila)-5 :]`.

  - FECHAS: ninguno de los 3 ficheros tiene fecha en el nombre. Metadatos
    docx.core_properties inspeccionados con python-docx sobre los 3:

        "REVISION EGURROLA .docx"   creado 02/01/2025 16:18  modificado 07/01/2025 10:23
        "REVISION EGURROLA 2.docx"  creado 02/01/2025 17:46  modificado 02/01/2025 17:46
        "REVISION EGURROLA 3.docx"  creado 02/01/2025 17:56  modificado 02/01/2025 17:56

    DECISION (documentada tal y como pide el procedimiento): se usa
    `created`, NUNCA `modified`/mtime, ni para ordenar ni para la fecha
    mostrada. Motivo, verificado celda a celda contra el CONTENIDO real
    (no es una suposicion):

      * Contando cuantas celdas de estado estan en 'X' o 'M' en cada
        fichero, el avance es estrictamente creciente en el ORDEN DE
        CREACION: "sin numero" (16:18) < "2" (17:46) < "3" (17:56). Por
        ejemplo, en planta 0: "sin numero" tiene Montante teleco/sscc en
        'M' y Tubeado/Cableado zzcc vacios; "2" ya tiene esos mismos items
        en 'X' y anade "1as caras Pladur"; "3" anade ademas "Cableado
        zzcc" en 'X'. Es coherente con 3 capturas del MISMO dia de trabajo
        (02/01/2025), guardadas cada vez que llegaba mas informacion esa
        tarde.
      * Si se usara `modified` en su lugar, "REVISION EGURROLA .docx"
        pasaria a fecharse el 07/01/2025 -- DESPUES de "2" y "3" (ambos
        02/01) -- pese a tener MENOS avance marcado que ambas. Eso
        implicaria un retroceso de obra sin ninguna nota ni evidencia que
        lo explique, lo cual no tiene sentido real y contradice el propio
        contenido. Se interpreta que el guardado del 07/01 fue un retoque
        no relacionado con el estado de obra (candidato mas probable: fue
        cuando se le anadieron las 2 columnas vacias extra que solo tiene
        este fichero, ya que "2" y "3" no las tienen y nunca se volvieron
        a modificar tras su creacion). Esta contradiccion se deja senalada
        aqui explicitamente, no se resuelve inventando una fecha.
      * Consecuencia practica: los 3 ficheros caen en el MISMO dia
        calendario (02/01/2025) una vez formateados a 'DD/MM/AAAA' (regla
        de fechado del procedimiento). Se mantienen las 3 revisiones como
        3 entradas SEPARADAS del historial (no se descartan 2 de las 3),
        ordenadas por su instante de creacion completo, porque representan
        avance real y distinto verificado por contenido; descartar alguna
        perderia informacion real sin necesidad. Efecto colateral conocido
        y asumido: el grafico "Evolucion del avance en el tiempo" del
        panel mostrara 3 puntos bajo la misma etiqueta "02/01/2025" (no es
        un fallo del adaptador, es fiel a que hubo 3 guardados ese dia).
        La revision mas reciente/completa para los KPIs del snapshot
        actual es, correctamente, "REVISION EGURROLA 3.docx".

  - VOCABULARIO DE ESTADO: inspeccionado celda a celda (no una muestra) el
    contenido de las 3 tablas completas. Solo aparecen dos simbolos mas la
    celda vacia: 'X' y 'M'. NO aparece ningun '/' en ninguno de los 3
    ficheros. No existe en la carpeta de la obra ningun documento de
    leyenda propio (se ha buscado "leyenda"/"gemini"/"prompt tabla" y no
    hay nada equivalente al de Mungia). Se asume la MISMA convencion que
    en el resto de obras del sistema (X=terminado, M=avanzado >50%,
    vacio=no iniciado) por ser la convencion compartida por Sagarde en
    todas las obras vistas hasta ahora, pero esto es una INFERENCIA por
    coherencia con el resto del proyecto, no una confirmacion documental
    propia de esta obra -- se deja dicho aqui y en el informe final.
    Cualquier simbolo no reconocido se deja pasar tal cual (no se pierde
    informacion) y se avisa por consola.

Devuelve un "historial": lista de (fecha_dd/mm/aaaa, snapshot) ordenada por
fecha, en el esquema estandar que espera motor_informes.py.
"""
import os
import docx

CARPETA_OBRA = os.path.join(
    os.path.dirname(os.path.dirname(os.path.dirname(os.path.abspath(__file__)))),
    "2025 GETXO 12V EGURROLA"
)

# Edificio unico confirmado (ningun documento de la obra menciona mas de un
# portal). Ver docstring: el numero de portal exacto no esta confirmado.
BUILDING = "Artibai"

# Ver docstring: confirmado contra los planos de "2025 12v egurrola getxo.pdf"
# (PLANTA BAJA / PLANTA PRIMERA / PLANTA SEGUNDA / PLANTA TERCERA).
FLOOR_LABELS = {'0': 'PB', '1': '1', '2': '2', '3': '3'}

_SIMBOLOS_DESCONOCIDOS = set()


def _normalizar_estado(tarea, v):
    v = (v or '').strip()
    if v in ('X', 'M', ''):
        return v
    _SIMBOLOS_DESCONOCIDOS.add((tarea, v))
    return v


def _parsear_tabla(d):
    """d: objeto docx.Document ya cargado (se reutiliza, ver cargar_historial)."""
    if not d.tables:
        return []
    filas = [[c.text.strip() for c in r.cells] for r in d.tables[0].rows]
    registros = []
    unit_labels_l = unit_labels_r = None

    for fila in filas:
        n = len(fila)
        if n < 10:
            continue  # fila residual/corta al final de la tabla, sin datos

        if fila[0] == 'TRABAJO' and fila[1] == 'PLANTA':
            # fila de cabecera de bloque: da los codigos reales de vivienda
            idx_r = n - 5
            unit_labels_l = fila[2:5]
            unit_labels_r = fila[idx_r + 2: idx_r + 5]
            continue

        if unit_labels_l is None:
            continue  # aun no hemos visto ninguna cabecera de bloque

        idx_r = n - 5
        task_l, floor_l = fila[0], fila[1]
        vals_l = fila[2:5]
        task_r, floor_r = fila[idx_r], fila[idx_r + 1]
        vals_r = fila[idx_r + 2: idx_r + 5]

        if task_l and floor_l:
            floor_disp = FLOOR_LABELS.get(floor_l, floor_l)
            for unit, val in zip(unit_labels_l, vals_l):
                if unit:
                    registros.append({
                        'task': task_l, 'floor': floor_disp,
                        'building': BUILDING, 'unit': unit,
                        'status': _normalizar_estado(task_l, val),
                    })
        if task_r and floor_r:
            floor_disp = FLOOR_LABELS.get(floor_r, floor_r)
            for unit, val in zip(unit_labels_r, vals_r):
                if unit:
                    registros.append({
                        'task': task_r, 'floor': floor_disp,
                        'building': BUILDING, 'unit': unit,
                        'status': _normalizar_estado(task_r, val),
                    })
    return registros


def cargar_historial():
    if not os.path.isdir(CARPETA_OBRA):
        raise FileNotFoundError(f"No se encuentra la carpeta de la obra: {CARPETA_OBRA}")

    candidatos = []
    for fn in os.listdir(CARPETA_OBRA):
        fn_low = fn.lower()
        if fn_low.startswith('revision egurrola') and fn_low.endswith('.docx'):
            candidatos.append(os.path.join(CARPETA_OBRA, fn))

    # Se carga cada documento UNA vez (se reutiliza tanto para leer la fecha
    # de creacion como para parsear la tabla). Ver docstring: se ordena y se
    # fecha por `core_properties.created`, nunca por mtime/modified.
    info = []
    for ruta in candidatos:
        d = docx.Document(ruta)
        creado = d.core_properties.created
        if creado is None:
            print(f"  [aviso] {os.path.basename(ruta)} no tiene fecha de creacion en metadatos, se omite")
            continue
        info.append((creado, os.path.basename(ruta), ruta, d))
    info.sort(key=lambda x: x[0])

    historial = []
    for creado, fn, ruta, d in info:
        display = creado.strftime('%d/%m/%Y')
        registros = _parsear_tabla(d)
        if registros:
            historial.append((display, registros))
        else:
            print(f"  [sin tablas de datos] {fn} no aporta registros, se omite")

    if _SIMBOLOS_DESCONOCIDOS:
        print(f"  [aviso] simbolos de estado no reconocidos encontrados: {sorted(_SIMBOLOS_DESCONOCIDOS)}")

    return historial


if __name__ == "__main__":
    import sys

    h = cargar_historial()
    print(f"\nRevisiones cargadas: {len(h)}")
    if h:
        print(f"Primera: {h[0][0]}  ({len(h[0][1])} registros)")
        print(f"Última:  {h[-1][0]}  ({len(h[-1][1])} registros)")

        SISTEMA_DIR = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
        sys.path.insert(0, SISTEMA_DIR)
        import motor_informes

        kpis = motor_informes.kpis_snapshot(h[-1][1])
        print(f"\nKPIs snapshot última revisión ({h[-1][0]}):")
        print(f"  {kpis}")

        bloqueos = motor_informes.detectar_bloqueos(h[-1][1])
        print(f"Bloqueos detectados: {len(bloqueos)}")
    else:
        print("No se ha cargado ninguna revisión con datos.")
