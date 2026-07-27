# -*- coding: utf-8 -*-
"""
ADAPTADOR - 2024 BILBAO 88V ZORROZAURE  (capa 1 - lectura, especifico de obra)
--------------------------------------------------------------------------------
Formato de esta obra (verificado con python-docx sobre el UNICO fichero de
revision que existe en la carpeta): "REVISION zorrozaure.docx" esta en la
RAIZ de la carpeta de la obra, no en una subcarpeta "REVISIONES" como en
Mungia/Boluета/Obispo Orueta. No hay serie temporal: es un solo snapshot,
asi que `cargar_historial()` devuelve una lista de UNA sola entrada (esto
es correcto, no un fallo del adaptador ni una revision que falte por
localizar).

ESTRUCTURA DE TABLAS (verificada celda a celda, las 4 tablas completas):
  - 4 tablas en total, 2 EDIFICIOS ("Bloque A1" y "Bloque A2"), cada uno
    repartido en 2 tablas de 68 filas:
      Tabla 0 = Bloque A1, plantas 1-4   (68 filas x 22 columnas)
      Tabla 1 = Bloque A1, plantas 5-8   (68 filas x 22 columnas)
      Tabla 2 = Bloque A2, plantas 1-4   (68 filas x 10 columnas)
      Tabla 3 = Bloque A2, plantas 5-8   (68 filas x 10 columnas)
  - Cada tabla trae, en la fila 0, el nombre del edificio en una celda
    fusionada ("Bloque A1" / "Bloque A2"). Dentro de la tabla hay 2
    "bloques" de 2 plantas en paralelo (izquierda = planta N, derecha =
    planta N+1), cada uno con su propia fila de cabecera de unidades
    (columnas 0 y 1 vacias, columna 2 = primera letra de unidad). Esta
    cabecera se detecta DINAMICAMENTE (no se asume una fila fija), igual
    que en el adaptador de Bolueta, por si una revision futura cambiara
    el numero de filas de tareas.
  - IMPORTANTE, y DISTINTO de lo apuntado antes de abrir el fichero: el
    Bloque A1 tiene 9 unidades por planta (A..I), pero el Bloque A2 tiene
    SOLO 3 unidades por planta (A, B, C) -- verificado en las 4 tablas
    (cabecera de unidades siempre 'A','B','C' en las tablas 2 y 3, nunca
    llega a 'I'). El aviso de "9 unidades" del prompt original solo era
    cierto para el Bloque A1.
  - Cada tabla cubre 8 plantas en total (2 tablas x 4 plantas cada una =
    8 plantas por edificio), numeradas '1'..'8' como texto.

VALORES DE ESTADO: se ha contado exhaustivamente TODAS las celdas de datos
de las 4 tablas (2577 celdas de datos). Unicos valores encontrados:
  'X' (1997), '' vacio (550), 'M' (15), '/' (15).
No aparecio ningun otro simbolo suelto en las celdas de datos reales; las
"letras/digitos sueltos" mencionadas de un barrido rapido previo eran, en
efecto, ruido de las filas de cabecera (letras de unidad A-I) y no de las
celdas de estado -- confirmado, no hacia falta remapeo alguno.

Ademas, estos 4 valores coinciden EXACTAMENTE con el diccionario estandar
ya definido dentro de motor_informes.py (SCORE / ESTADO_LABEL: 'X'=
terminado 100%, 'M'=avanzado >50%, '/'=iniciado <50%, ''=no iniciado), que
es la misma convencion usada en Mungia y Boluета. No se ha encontrado en
la carpeta de esta obra ningun documento de "leyenda" propio (se buscaron
nombres tipo leyenda/gemini/prompt/clave/simbolo: no existe ninguno;
"Checklist_Tareas_Bloques_A1_A2.docx" es un parte de incidencias de
ascensores e iluminacion de zonas comunes, NO una leyenda de la tabla de
revision). Se asume el diccionario estandar del motor por coincidir de
forma exacta con los valores observados, no por inferencia arriesgada.

FECHA DE LA REVISION: el fichero no tiene fecha en el nombre. Se usa la
fecha que aparece en la CABECERA DE PAGINA interna del propio documento
Word ("ZORROZAURE 06/09/2024   SEMANA : 28"), por ser contenido redactado
explicitamente por quien hizo la revision (mas fiable que un metadato de
sistema que puede alterarse en copias/sincronizaciones). Se ha cruzado
con dos metadatos independientes y los tres COINCIDEN:
  - Cabecera interna del documento:      06/09/2024
  - mtime del fichero en disco:          06/09/2024 07:52:47 (hora local)
  - docx.core_properties.modified:       06/09/2024 05:52:00 UTC
      (05:52 UTC == 07:52 CEST, mismo instante que el mtime)
Se implementa con fallback automatico a mtime si en el futuro se
sustituye este fichero por otro sin cabecera interna reconocible.

Devuelve un "historial": lista de (fecha_dd/mm/aaaa, snapshot) en el
esquema estandar que espera motor_informes.py. Si el formato de esta obra
cambia (nueva revision, tareas distintas), este es el UNICO fichero que
hay que tocar. El motor_informes.py no se toca nunca.
"""
import os
import re
import datetime
import docx

RUTA_OBRA = os.path.join(
    os.path.dirname(os.path.dirname(os.path.dirname(os.path.abspath(__file__)))),
    "2024 BILBAO 88V ZORROZAURE"
)
RUTA_REVISION = os.path.join(RUTA_OBRA, "REVISION zorrozaure.docx")

RE_CABECERA_FECHA = re.compile(r'(\d{2})/(\d{2})/(\d{4})')
RE_LETRA_UNIDAD = re.compile(r'^[A-Z]$')


def _fecha_desde_documento(d, ruta):
    """Ver docstring del modulo: prioridad = cabecera interna del docx,
    con fallback a mtime del fichero si no se encuentra o no hay match."""
    try:
        for sec in d.sections:
            for p in sec.header.paragraphs:
                m = RE_CABECERA_FECHA.search(p.text or '')
                if m:
                    return f"{m.group(1)}/{m.group(2)}/{m.group(3)}"
    except Exception:
        pass
    mtime = os.path.getmtime(ruta)
    return datetime.datetime.fromtimestamp(mtime).strftime('%d/%m/%Y')


def _es_fila_cabecera_unidades(row, ncols):
    if len(row) < ncols or ncols < 3:
        return False
    return row[0] == '' and row[1] == '' and bool(row[2]) and bool(RE_LETRA_UNIDAD.match(row[2]))


def _parsear_tabla(tabla):
    rows = [[c.text.strip() for c in r.cells] for r in tabla.rows]
    n = len(rows)
    if n < 2:
        return []
    ncols = len(tabla.columns)
    mid = ncols // 2  # 11 para Bloque A1 (22 cols), 5 para Bloque A2 (10 cols)

    building = (rows[0][0] or '').strip() if rows[0] else ''
    if not building:
        building = '[Edificio no identificado en fila de titulo]'

    registros = []
    i = 0
    while i < n:
        row = rows[i]
        if not _es_fila_cabecera_unidades(row, ncols):
            i += 1
            continue
        header = row
        j = i + 1
        while j < n:
            r = rows[j]
            if len(r) < ncols:
                j += 1
                continue
            if _es_fila_cabecera_unidades(r, ncols):
                break  # siguiente bloque de 2 plantas
            task_l, floor_l = r[0].strip(), r[1].strip()
            task_r, floor_r = r[mid].strip(), r[mid + 1].strip()
            if task_l and floor_l:
                for col in range(2, mid):
                    unidad = header[col] if col < len(header) else ''
                    if unidad:
                        registros.append({
                            'task': task_l, 'floor': floor_l,
                            'building': building, 'unit': unidad,
                            'status': r[col] if col < len(r) else '',
                        })
            if task_r and floor_r:
                for col in range(mid + 2, ncols):
                    unidad = header[col] if col < len(header) else ''
                    if unidad:
                        registros.append({
                            'task': task_r, 'floor': floor_r,
                            'building': building, 'unit': unidad,
                            'status': r[col] if col < len(r) else '',
                        })
            j += 1
        i = j
    return registros


def cargar_historial():
    if not os.path.isfile(RUTA_REVISION):
        raise FileNotFoundError(f"No se encuentra el fichero de revision: {RUTA_REVISION}")

    d = docx.Document(RUTA_REVISION)
    display = _fecha_desde_documento(d, RUTA_REVISION)

    registros = []
    for tabla in d.tables:
        registros.extend(_parsear_tabla(tabla))

    historial = []
    if registros:
        historial.append((display, registros))
    else:
        print(f"  [sin registros] {os.path.basename(RUTA_REVISION)} no aporto datos de tabla")
    return historial


if __name__ == "__main__":
    import sys

    h = cargar_historial()
    print(f"Revisiones cargadas: {len(h)}  (se espera 1: solo hay un snapshot, sin serie temporal)")
    if h:
        fecha, registros = h[0]
        print(f"Fecha: {fecha}  ({len(registros)} registros)")

        edificios = sorted(set(r['building'] for r in registros))
        plantas = sorted(set(r['floor'] for r in registros), key=lambda x: int(x) if x.isdigit() else 99)
        tareas = sorted(set(r['task'] for r in registros))
        estados = {}
        for r in registros:
            estados[r['status']] = estados.get(r['status'], 0) + 1
        print(f"Edificios: {edificios}")
        print(f"Plantas: {plantas}")
        print(f"Nº tareas distintas: {len(tareas)}")
        print(f"Conteo de estados: {estados}")

        for ed in edificios:
            unidades = sorted(set(r['unit'] for r in registros if r['building'] == ed))
            print(f"  Unidades en {ed}: {unidades}")

        SISTEMA_DIR = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
        sys.path.insert(0, SISTEMA_DIR)
        import motor_informes

        kpis = motor_informes.kpis_snapshot(registros)
        print(f"\nKPIs snapshot ({fecha}):")
        print(f"  {kpis}")

        bloqueos = motor_informes.detectar_bloqueos(registros)
        print(f"Bloqueos detectados: {len(bloqueos)}")
    else:
        print("No se ha cargado ninguna revisión con datos.")
