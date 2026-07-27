#!/usr/bin/env python3
"""
postventas_sync.py — Sincronizador de incidencias Post-Ventas Sagarde
======================================================================
Detecta el Word matriz y los PDFs de incidencias resueltas en cada
carpeta. Cruza los datos y añade al Word las incidencias que falten,
manteniendo el formato original al 100%.

Solo añade automáticamente los registros de PDFs DIGITALES (tipados),
donde el campo "CÓDIGO SAGARDE" es legible con precisión. Los PDFs
manuscritos/escaneados se marcan para revisión manual.

Uso:
    python3 postventas_sync.py                           # todas las carpetas
    python3 postventas_sync.py --carpeta DINAMITA        # solo esa
    python3 postventas_sync.py --tecnico Manuel --fecha 15/07/2026
    python3 postventas_sync.py --dry-run                 # solo muestra
"""

import os, re, copy, shutil, argparse, sys
from pathlib import Path
from datetime import date

try:
    import pdfplumber
except ImportError:
    sys.exit("pip install pdfplumber --break-system-packages")
try:
    from docx import Document
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
except ImportError:
    sys.exit("pip install python-docx --break-system-packages")

# ── Patrones ──────────────────────────────────────────────────────────────────

# Código Sagarde: 2-6 letras + 9-12 dígitos  (JFAI20251100361, MLRI2026501606…)
CODE_RE = re.compile(r'\b[A-Za-z]{2,6}\d{9,12}\b')

# Campo explícito en PDFs digitales — ÚNICA fuente fiable
CODIGO_LABEL_RE = re.compile(
    r'(?:C[ÓO]DIGO\s+SAGARDE|COD[IÍ]GO\s+SAGARDE)[:\s]+([A-Za-z]{2,6}\d{9,12})',
    re.I
)

# Resto de campos del PDF digital
FIELDS = {
    "portal":   re.compile(r'PORTAL[:\s]+[A-Za-z]*\s*(\w+)', re.I),
    "mano":     re.compile(r'(?:PISO\s*/\s*MANO|MANO)[:\s]+[Pp]uerta\s+([A-Za-z])', re.I),
    "cliente":  re.compile(r'CLIENTE[:\s]+(.+?)(?:\n|TELÉFONO|TELEFONO)', re.I),
    "telefono": re.compile(r'(?:TELÉFONO|TELEFONO)[^:\n]*:\s*([0-9 /\-]+)', re.I),
    "fecha_av": re.compile(r'FECHA\s+AVISO[:\s]+([^\n]+)', re.I),
    "fecha_re": re.compile(r'FECHA\s+DE\s+REAL[IÍ]Z[AÁ]CI[ÓO]N[:\s]+([^\n]+)', re.I),
    "desc":     re.compile(
        r'DESCRIPCI[ÓO]N\s+DE\s+LA\s+INCIDENCIA[:\s]*\n(.+?)(?:OBSERVACIONES|FIRMA|$)',
        re.I | re.S
    ),
}


# ── Extracción PDF ────────────────────────────────────────────────────────────

def extraer_pagina(text: str, pagina: int, pdf_name: str) -> dict:
    """
    Extrae datos de una página de PDF.
    Solo marca como fiable si el código aparece con la etiqueta explícita
    'CÓDIGO SAGARDE:'. Páginas sin esa etiqueta = manuscritas = ilegibles.
    """
    base = {"pagina": pagina, "pdf": pdf_name}

    m = CODIGO_LABEL_RE.search(text)
    if not m:
        # PDF manuscrito o sin estructura legible → revisión manual
        return {**base, "codigo": None, "ilegible": True, "texto_raw": text[:200]}

    data = {**base, "codigo": m.group(1).strip(), "ilegible": False}

    for campo, pat in FIELDS.items():
        m2 = pat.search(text)
        data[campo] = m2.group(1).strip()[:150] if m2 else ""

    if data.get("desc"):
        data["desc"] = " ".join(data["desc"].split())[:250]

    return data


def leer_pdf(pdf_path: Path) -> list[dict]:
    resultados = []
    try:
        with pdfplumber.open(str(pdf_path)) as pdf:
            for i, page in enumerate(pdf.pages, 1):
                text = page.extract_text() or ""
                resultados.append(extraer_pagina(text, i, pdf_path.name))
    except Exception as e:
        print(f"    ⚠  Error: {pdf_path.name}: {e}")
    return resultados


# ── Word ──────────────────────────────────────────────────────────────────────

def codigos_en_word(table) -> set:
    """Códigos Sagarde presentes en la columna 4 de la tabla."""
    codes = set()
    for row in table.rows:
        codes.update(CODE_RE.findall(row.cells[4].text))
    return codes


def detectar_formato(table) -> tuple[str, str]:
    """Devuelve (font_name, sz_half_points) del primer run encontrado."""
    for row in table.rows:
        for cell in row.cells:
            for para in cell.paragraphs:
                for run in para.runs:
                    fn = run.font.name or "Arial Narrow"
                    sz = str(int(run.font.size // 6350)) if run.font.size else "18"
                    return fn, sz
    return "Arial Narrow", "18"


def build_rPr(fn: str, sz: str):
    rPr = OxmlElement('w:rPr')
    rf = OxmlElement('w:rFonts')
    rf.set(qn('w:ascii'), fn); rf.set(qn('w:hAnsi'), fn)
    rPr.append(rf)
    for tag in ('w:sz', 'w:szCs'):
        el = OxmlElement(tag); el.set(qn('w:val'), sz); rPr.append(el)
    return rPr


def añadir_fila(table, datos: list, fn: str, sz: str):
    src_row = table.rows[-1]
    new_tr  = copy.deepcopy(src_row._tr)
    src_cells = src_row._tr.findall(qn('w:tc'))

    for i, cell_elem in enumerate(new_tr.findall(qn('w:tc'))):
        if i >= len(datos):
            break
        for p in cell_elem.findall(qn('w:p')):
            cell_elem.remove(p)
        texto  = str(datos[i]) if datos[i] else ""
        lineas = texto.split('\n') if texto else ['']
        ref_p  = (src_cells[i].findall(qn('w:p')) or [None])[0]

        for linea in lineas:
            new_p = copy.deepcopy(ref_p) if ref_p is not None else OxmlElement('w:p')
            for r in new_p.findall(qn('w:r')):
                new_p.remove(r)
            if linea:
                run = OxmlElement('w:r')
                run.append(build_rPr(fn, sz))
                t = OxmlElement('w:t')
                t.text = linea
                t.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
                run.append(t)
                new_p.append(run)
            cell_elem.append(new_p)

    table._tbl.append(new_tr)


def ordenar_tabla(table):
    tbl = table._tbl
    trs = tbl.findall(qn('w:tr'))
    def clave(tr):
        cells = tr.findall(qn('w:tc'))
        def txt(c): return ''.join(x.text or '' for x in c.iter(qn('w:t'))).strip()
        return (txt(cells[1]).zfill(6), txt(cells[2]))
    for tr in trs:
        tbl.remove(tr)
    for tr in sorted(trs, key=clave):
        tbl.append(tr)


def fecha_corta(f: str) -> str:
    p = f.strip().split('/')
    return f"{p[0]}/{p[1]}/{p[2][2:]}" if len(p) == 3 and len(p[2]) == 4 else f


# ── Detección de archivos ─────────────────────────────────────────────────────

def buscar_word(folder: Path) -> Path | None:
    docxs = [f for f in folder.glob("*.docx") if "backup" not in f.name.lower()]
    for f in docxs:
        if "postventa" in f.name.lower():
            return f
    if len(docxs) == 1:
        return docxs[0]
    for f in docxs:
        try:
            doc = Document(str(f))
            if any(len(t.columns) >= 9 for t in doc.tables):
                return f
        except:
            pass
    return None


def buscar_pdfs(folder: Path) -> list[Path]:
    kw = ["solucionada", "resuelta", "incidencia"]
    return sorted(f for f in folder.glob("*.pdf")
                  if any(k in f.name.lower() for k in kw))


# ── Procesado de carpeta ──────────────────────────────────────────────────────

def procesar(folder: Path, tecnico: str, fecha: str, dry_run: bool) -> dict:
    r = dict(carpeta=folder.name, word=None, num_pdfs=0,
             codigos_word=0, nuevas=[], ilegibles=[], error=None)

    wp = buscar_word(folder)
    if not wp:
        r["error"] = "Sin Word matriz"; return r
    r["word"] = wp.name

    pdfs = buscar_pdfs(folder)
    r["num_pdfs"] = len(pdfs)
    if not pdfs:
        r["error"] = "Sin PDFs de incidencias"; return r

    doc = Document(str(wp))
    if not doc.tables:
        r["error"] = "Word sin tablas"; return r
    table = doc.tables[0]
    if len(table.columns) < 9:
        r["error"] = f"Tabla con {len(table.columns)} cols (esperadas ≥9)"; return r

    existing  = codigos_en_word(table)
    r["codigos_word"] = len(existing)
    fn, sz    = detectar_formato(table)

    nuevas = {}
    for pdf_path in pdfs:
        for inc in leer_pdf(pdf_path):
            if inc.get("ilegible"):
                r["ilegibles"].append(f"{inc['pdf']} p.{inc['pagina']}")
            elif inc["codigo"] not in existing:
                if inc["codigo"] not in nuevas:
                    nuevas[inc["codigo"]] = inc

    r["nuevas"] = list(nuevas.keys())

    if not dry_run and nuevas:
        bk = wp.with_name(wp.stem + "_BACKUP" + wp.suffix)
        if not bk.exists():
            shutil.copy(str(wp), str(bk))

        for cod, inc in nuevas.items():
            cliente  = inc.get("cliente", "").strip()
            tel      = inc.get("telefono", "").strip()
            cel0     = f"{cliente}\n{tel}".strip() if tel else cliente
            desc     = inc.get("desc", "").strip()
            col4     = f"{cod}\n{desc}" if desc else cod

            añadir_fila(table, [
                cel0,                   # cliente
                inc.get("portal",""),   # portal
                inc.get("mano",""),     # mano
                "",                     # ref
                col4,                   # código + descripción
                inc.get("fecha_av",""), # fecha aviso
                tecnico,                # técnico
                fecha,                  # fecha resolución
                "Si",                   # resuelta
                fecha_corta(fecha),     # fecha corta
            ], fn, sz)

        ordenar_tabla(table)
        doc.save(str(wp))

    return r


# ── Main ──────────────────────────────────────────────────────────────────────

def main():
    ap = argparse.ArgumentParser()
    ap.add_argument("--carpeta",  default="")
    ap.add_argument("--tecnico",  default="Manuel")
    ap.add_argument("--fecha",    default="")
    ap.add_argument("--dry-run",  action="store_true")
    args = ap.parse_args()

    if not args.fecha:
        args.fecha = date.today().strftime("%d/%m/%Y")

    root = Path(__file__).parent
    dirs = sorted(d for d in root.iterdir()
                  if d.is_dir() and d.name.upper().startswith("INCIDENCIAS"))
    if args.carpeta:
        dirs = [d for d in dirs if args.carpeta.upper() in d.name.upper()]

    modo = "DRY-RUN" if args.dry_run else "ACTUALIZANDO"
    print(f"\n{'═'*65}")
    print(f"  POST-VENTAS SAGARDE  ·  {modo}")
    print(f"  Técnico: {args.tecnico}  ·  Fecha: {args.fecha}")
    print(f"{'═'*65}\n")

    total = 0
    for d in dirs:
        r = procesar(d, args.tecnico, args.fecha, args.dry_run)
        ico = "⚠ " if r["error"] else ("✅" if r["nuevas"] else "✓ ")
        print(f"{ico} {r['carpeta']}")
        if r["error"]:
            print(f"     → {r['error']}")
        else:
            print(f"     Word: {r['word']}  "
                  f"({r['codigos_word']} registros, {r['num_pdfs']} PDF(s))")
            if r["nuevas"]:
                accion = "para añadir" if args.dry_run else "añadidas"
                print(f"     Incidencias {accion} ({len(r['nuevas'])}):")
                for c in r["nuevas"]:
                    print(f"       + {c}")
                total += len(r["nuevas"])
            else:
                print(f"     ✓ Todo al día")
            if r["ilegibles"]:
                n = len(r["ilegibles"])
                print(f"     ⚠  {n} página(s) manuscrita(s) — revisar manualmente")
        print()

    print(f"{'─'*65}")
    accion = "encontradas" if args.dry_run else "añadidas"
    print(f"  Incidencias nuevas {accion}: {total}\n")

if __name__ == "__main__":
    main()
