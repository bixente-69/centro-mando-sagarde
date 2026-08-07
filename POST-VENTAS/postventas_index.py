#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Genera index.html para la carpeta POST-VENTAS.

La app muestra una tarjeta por carpeta de incidencias, ordenada por el
archivo mas reciente encontrado dentro de cada obra.
"""
from __future__ import annotations

from datetime import datetime
from hashlib import sha1
from html import escape
from pathlib import Path
import json
import re
from urllib.parse import quote


ROOT = Path(__file__).resolve().parent
INDEX_PATH = ROOT / "index.html"
RESUMEN_JSON = ROOT / "postventas_resumen.json"
PREVIEW_DIR_NAME = "_PREVIEWS_WORD"
PREVIEW_DIR = ROOT / PREVIEW_DIR_NAME

IGNORE_DIRS = {".memory", "__pycache__", PREVIEW_DIR_NAME, "_SISTEMA"}
IGNORE_FILES = {"desktop.ini", "thumbs.db", "index.html"}
IMAGE_EXTS = {".jpg", ".jpeg", ".png", ".gif", ".bmp", ".heic", ".webp"}
WORD_EXTS = {".doc", ".docx"}
EXCEL_EXTS = {".xls", ".xlsx", ".xlsm"}


def fmt_dt(ts: float | None) -> str:
    if not ts:
        return "-"
    return datetime.fromtimestamp(ts).strftime("%d/%m/%Y %H:%M")


def days_text(ts: float | None) -> str:
    if not ts:
        return "sin fecha"
    d = (datetime.now().date() - datetime.fromtimestamp(ts).date()).days
    if d <= 0:
        return "hoy"
    if d == 1:
        return "ayer"
    return f"hace {d} días"


def url_for(path: Path) -> str:
    rel = path.relative_to(ROOT).as_posix()
    return quote(rel, safe="/()[]!$&'()*+,;=:@-._~")


def clean_name(folder_name: str) -> str:
    name = " ".join(folder_name.split())
    if name.upper().startswith("INCIDENCIAS "):
        name = name[11:].strip()
    return name or folder_name


def iter_files(folder: Path) -> list[Path]:
    files: list[Path] = []
    for path in folder.rglob("*"):
        if path.is_dir():
            continue
        if any(part in IGNORE_DIRS for part in path.relative_to(ROOT).parts):
            continue
        if path.name.startswith("~$") or path.name.lower() in IGNORE_FILES:
            continue
        files.append(path)
    return files


def latest(paths: list[Path]) -> Path | None:
    if not paths:
        return None
    return max(paths, key=lambda p: p.stat().st_mtime)


def primary_word(words: list[Path]) -> Path | None:
    valid = [p for p in words if "backup" not in p.name.lower()]
    preferred = [
        p for p in valid
        if "postventa" in p.name.lower() or "incidencia" in p.name.lower()
    ]
    return latest(preferred or valid or words)


def scan_obras() -> list[dict]:
    obras: list[dict] = []
    for folder in ROOT.iterdir():
        if not folder.is_dir() or folder.name in IGNORE_DIRS:
            continue
        if folder.name.startswith("_"):
            continue
        files = iter_files(folder)
        pdfs = [p for p in files if p.suffix.lower() == ".pdf"]
        words = [p for p in files if p.suffix.lower() in WORD_EXTS]
        excels = [p for p in files if p.suffix.lower() in EXCEL_EXTS]
        images = [p for p in files if p.suffix.lower() in IMAGE_EXTS]
        last_file = latest(files)
        last_ts = last_file.stat().st_mtime if last_file else folder.stat().st_mtime
        latest_files = sorted(files, key=lambda p: p.stat().st_mtime, reverse=True)[:4]
        main_word = primary_word(words)

        obras.append({
            "folder": folder,
            "name": clean_name(folder.name),
            "raw_name": folder.name,
            "files": files,
            "pdfs": pdfs,
            "words": words,
            "excels": excels,
            "images": images,
            "main_word": main_word,
            "last_pdf": latest(pdfs),
            "last_file": last_file,
            "last_ts": last_ts,
            "last_incident_ts": ultima_fecha_incidencia(main_word),
            "latest_files": latest_files,
        })

    # Un recuento de 0 es senal de alarma, no de "no aplica". Si este script
    # se mueve de carpeta, ROOT deja de apuntar a POST-VENTAS y el indice
    # saldria vacio con codigo de salida 0: el .bat solo mira errorlevel.
    if not obras:
        raise SystemExit(
            f"[ERROR] Ninguna carpeta de incidencias bajo {ROOT}. "
            f"Si el script se ha movido, ROOT esta mal calculado. "
            f"No se reescribe index.html con un indice vacio.")

    return sorted(obras, key=lambda o: (-o["last_ts"], o["name"].casefold()))


def stat_block(label: str, value: int | str, cls: str = "") -> str:
    cls_attr = f" {cls}".rstrip()
    return f'<div class="stat"><div class="stat-value{escape(cls_attr, quote=True)}">{escape(str(value))}</div><div class="stat-label">{escape(label)}</div></div>'


def action_link(label: str, path: Path | None, class_name: str = "") -> str:
    if not path:
        return f'<span class="btn disabled">{escape(label)}</span>'
    cls = f"btn {class_name}".strip()
    return f'<a class="{cls}" href="{escape(url_for(path), quote=True)}">{escape(label)}</a>'


def preview_path_for(docx_path: Path) -> Path:
    rel = docx_path.relative_to(ROOT).as_posix()
    digest = sha1(rel.encode("utf-8")).hexdigest()[:10]
    stem = re.sub(r"[^A-Za-z0-9._-]+", "_", docx_path.stem).strip("_")[:80] or "word"
    return PREVIEW_DIR / f"{stem}_{digest}.html"


def iter_docx_blocks(document):
    from docx.document import Document as DocxDocument
    from docx.table import Table
    from docx.text.paragraph import Paragraph

    if isinstance(document, DocxDocument):
        parent_elm = document.element.body
    else:
        parent_elm = document._tc

    for child in parent_elm.iterchildren():
        if child.tag.endswith("}p"):
            yield Paragraph(child, document)
        elif child.tag.endswith("}tbl"):
            yield Table(child, document)


def render_paragraph(paragraph) -> str:
    parts = []
    for run in paragraph.runs:
        text = escape(run.text).replace("\n", "<br>")
        if not text:
            continue
        if run.bold:
            text = f"<strong>{text}</strong>"
        if run.italic:
            text = f"<em>{text}</em>"
        if run.underline:
            text = f"<u>{text}</u>"
        parts.append(text)

    text = "".join(parts).strip()
    if not text:
        return ""

    style = (paragraph.style.name if paragraph.style else "").lower()
    if "title" in style or "título" in style:
        return f"<h2>{text}</h2>"
    if "heading 1" in style or "encabezado 1" in style:
        return f"<h2>{text}</h2>"
    if "heading" in style or "encabezado" in style:
        return f"<h3>{text}</h3>"
    return f"<p>{text}</p>"


def render_cell(cell) -> str:
    lines = []
    for paragraph in cell.paragraphs:
        text = " ".join(paragraph.text.split())
        if text:
            lines.append(escape(text))
    return "<br>".join(lines) or "&nbsp;"


# Formato estandar de Sagarde: Cliente, Portal, Mano, Ref, Codigo+Desc,
# FechaAviso, Tecnico, [TecnicoSecundario], Solucion, Resuelta, FechaCorta.
# La tabla tiene 10 columnas normalmente, pero algunas obras usan 11 (con un
# tecnico secundario extra). La columna "Resuelta" es siempre la penultima
# celda de la fila, no un indice fijo -- calcularla desde el final evita
# mirar la columna equivocada en las tablas de 11 columnas.
MIN_COLS_INCIDENCIAS = 9
RESOLVED_TOKENS = {"si", "s", "sí", "ok"}


def row_is_pending(row) -> bool:
    """Una fila se considera pendiente solo si:
    1) tiene el formato de tabla de incidencias (>=9 columnas),
    2) tiene contenido real (cliente o codigo/descripcion no vacios -- si no,
       es una fila de relleno vacia, no una incidencia), y
    3) su columna 'Resuelta' (penultima celda) no contiene un marcador de
       resuelto (Si/S/SI/Ok, sea cual sea el convenio de esa obra)."""
    cells = row.cells
    n = len(cells)
    if n < MIN_COLS_INCIDENCIAS:
        return False
    if not cells[0].text.strip() and not cells[4].text.strip():
        return False
    resuelta = cells[n - 2].text.strip().lower()
    return resuelta not in RESOLVED_TOKENS


EMBEDDED_DATE_RE = re.compile(r'\b(\d{1,2}/\d{1,2}/\d{2,4})\b')


def _parse_date_token(s: str):
    s = s.strip()
    for fmt in ("%d/%m/%Y", "%d/%m/%y"):
        try:
            return datetime.strptime(s, fmt)
        except ValueError:
            pass
    return None


def ultima_fecha_incidencia(word_path) -> float | None:
    """Fecha (timestamp) de la incidencia mas reciente registrada dentro de
    la tabla de incidencias, en vez de la fecha de modificacion del archivo
    .docx (que puede ser reciente solo porque se cerro/edito una incidencia
    vieja). Usa siempre la columna 5 ("fecha aviso" en el esquema estandar
    de >=9 columnas) -- nunca la penultima/ultima columna (resuelta / fecha
    de resolucion), porque esas SI se sobreescriben al cerrar una incidencia
    y contaminarian la recencia con la fecha del cierre en vez de la fecha
    real del aviso. Si la celda de fecha aviso esta vacia, intenta extraer
    una fecha del propio texto de la descripcion -- algunas obras (p.ej.
    MUNGUIA LARRABIZKER 94V) dejan esa columna vacia pero incluyen la fecha
    real al principio del texto. Ignora fechas en el futuro (errores de
    tecleo tipo "3018" en vez de "2018"). Solo aplica al esquema estandar
    (>=9 columnas); las tablas con esquema distinto (p.ej. GARELLANO, 8
    columnas sin fecha de aviso separada de la de resolucion) no se pueden
    leer de forma fiable aqui y caen al fallback por fecha de archivo."""
    if not word_path:
        return None
    try:
        from docx import Document
        doc = Document(str(word_path))
    except Exception:
        return None
    tables = [t for t in doc.tables if len(t.columns) >= MIN_COLS_INCIDENCIAS]
    if not tables:
        return None
    table = max(tables, key=lambda t: len(t.columns) * len(t.rows))
    ahora = datetime.now()
    dates = []
    for row in table.rows:
        cells = row.cells
        if len(cells) < 6:
            continue
        dt = _parse_date_token(cells[5].text.strip().split("\n")[0])
        if dt is None and len(cells) > 4:
            m = EMBEDDED_DATE_RE.search(cells[4].text)
            if m:
                dt = _parse_date_token(m.group(1))
        if dt and dt <= ahora:
            dates.append(dt)
    if not dates:
        return None
    return max(dates).timestamp()


# Obras cuyo esquema de tabla no es el estandar (por eso no se les puede
# calcular `last_incident_ts` de forma fiable, p.ej. porque la unica columna
# de fecha se sobreescribe con la fecha del dia al cerrar incidencias) pero
# que se sabe con certeza que llevan mas de 2 años sin actividad real y su
# postventa ya vencio. Añadir aqui el nombre limpio (`clean_name`) de la
# obra si aparece "Reciente" o sin "VENCIDO" por error debido a esta
# limitacion tecnica.
OBRAS_VENCIDAS_MANUAL = {"72V GARELLANO"}

POSTVENTA_WINDOW_DAYS = 2 * 365  # postventa = 2 años desde la entrega


def is_recent(obra: dict, window_days: int = 45) -> bool:
    """Recencia basada en la incidencia real mas reciente (si se pudo leer
    del Word) en vez de en la fecha de modificacion del archivo -- para que
    cerrar una incidencia antigua no la haga aparecer como 'Reciente'."""
    if obra["name"] in OBRAS_VENCIDAS_MANUAL:
        return False
    ts = obra.get("last_incident_ts")
    if ts is None:
        ts = obra["last_ts"]
    return (datetime.now().timestamp() - ts) <= window_days * 24 * 60 * 60


def is_vencido(obra: dict) -> bool:
    """La postventa de una obra dura 2 años desde la entrega. Si la ultima
    incidencia real registrada ya tiene mas de 2 años, la obra esta fuera de
    plazo: lo que se registre ahi a partir de ahora son casos aislados /
    gestion administrativa, ya no postventa propiamente dicha.

    Si no se puede leer una fecha de incidencia fiable (obras sin Word
    matriz como AMAZON/CHAO BERANGO, que no son registros de incidencias
    como GARDOKI/URDULIZ, o con fechas en formato no reconocido como
    PARKEDER/RESIDENCIA SANTUTXU), se usa la fecha de modificacion del
    archivo como respaldo -- igual que hace is_recent() -- porque si no se
    ha tocado hoy esa fecha sigue siendo la mejor senal disponible de su
    antiguedad real."""
    if obra["name"] in OBRAS_VENCIDAS_MANUAL:
        return True
    ts = obra.get("last_incident_ts")
    if ts is None:
        ts = obra["last_ts"]
    return (datetime.now().timestamp() - ts) > POSTVENTA_WINDOW_DAYS * 24 * 60 * 60


def render_table(table) -> tuple[str, int, int]:
    rows = []
    pending_count = 0
    pending_capable = len(table.columns) >= MIN_COLS_INCIDENCIAS
    for row in table.rows:
        pending = pending_capable and row_is_pending(row)
        if pending:
            pending_count += 1
        cells = "".join(f"<td>{render_cell(cell)}</td>" for cell in row.cells)
        cls = ' class="pending"' if pending else ""
        rows.append(f"<tr{cls}>{cells}</tr>")
    html = f'<div class="table-wrap"><table>{ "".join(rows) }</table></div>'
    return html, pending_count, (len(table.rows) if pending_capable else 0)


def write_preview(docx_path: Path) -> tuple[Path | None, int, int]:
    if docx_path.suffix.lower() != ".docx":
        return None, 0, 0

    preview_path = preview_path_for(docx_path)
    total_pending = 0
    total_rows = 0
    try:
        from docx import Document

        doc = Document(str(docx_path))
        body = []
        for block in iter_docx_blocks(doc):
            if hasattr(block, "rows"):
                table_html, pending_count, rows_considered = render_table(block)
                body.append(table_html)
                total_pending += pending_count
                total_rows += rows_considered
            else:
                html = render_paragraph(block)
                if html:
                    body.append(html)
        content = "\n".join(body) or '<p class="empty">Documento sin contenido de texto extraible.</p>'
        error = ""
        if total_pending:
            plural = "s" if total_pending != 1 else ""
            error = (
                f'<div class="banner pending">⚠ {total_pending} incidencia{plural} pendiente{plural} '
                f'de resolver (de {total_rows} registradas)</div>'
            )
    except Exception as exc:
        content = ""
        error = f'<div class="banner bad">No se pudo generar la vista previa: {escape(str(exc))}</div>'

    PREVIEW_DIR.mkdir(exist_ok=True)
    generated = datetime.now().strftime("%d/%m/%Y %H:%M")
    original_href = "../" + url_for(docx_path)
    index_href = "../index.html"
    html = f"""<!DOCTYPE html>
<html lang="es">
<head>
<meta charset="UTF-8">
<meta name="viewport" content="width=device-width, initial-scale=1.0">
<title>{escape(docx_path.name)}</title>
<style>
:root{{--bg:#f4f6f9;--card:#fff;--header:#0b1f3a;--header2:#123a63;--text:#1c2733;--muted:#647184;--accent:#f5a524;--bad:#d9483c;--pending:#e2680a;--line:#e3e7ee;--radius:10px;}}
*{{box-sizing:border-box;}}
body{{margin:0;font-family:-apple-system,BlinkMacSystemFont,'Segoe UI',Roboto,sans-serif;background:var(--bg);color:var(--text);line-height:1.45;}}
.wrap{{max-width:1180px;margin:0 auto;padding:20px;}}
.header{{background:linear-gradient(120deg,var(--header),var(--header2));color:#fff;border-radius:var(--radius);padding:20px 24px;margin-bottom:14px;display:flex;justify-content:space-between;gap:14px;align-items:flex-start;}}
.brand{{font-size:11px;letter-spacing:1.5px;text-transform:uppercase;color:var(--accent);font-weight:800;margin-bottom:4px;}}
h1{{font-size:20px;margin:0;line-height:1.25;}}
.meta{{font-size:12px;color:#c7d3e3;text-align:right;min-width:170px;}}
.actions{{display:flex;gap:7px;flex-wrap:wrap;margin-top:8px;justify-content:flex-end;}}
.btn{{display:inline-flex;align-items:center;justify-content:center;text-decoration:none;border:1px solid rgba(255,255,255,.35);color:#fff;border-radius:7px;padding:6px 10px;font-size:12px;font-weight:750;}}
.doc{{background:var(--card);border-radius:var(--radius);box-shadow:0 1px 3px rgba(0,0,0,.08);padding:22px;overflow:hidden;}}
.doc h2{{font-size:18px;margin:12px 0 8px;}}
.doc h3{{font-size:15px;margin:12px 0 7px;}}
.doc p{{font-size:13px;margin:7px 0;}}
.table-wrap{{overflow:auto;margin:12px 0;border:1px solid var(--line);border-radius:8px;}}
table{{width:100%;border-collapse:collapse;font-size:12px;background:#fff;}}
td{{border:1px solid var(--line);padding:6px 7px;vertical-align:top;min-width:72px;}}
tr:first-child td{{background:#f8f9fb;font-weight:750;}}
tr.pending td{{background:#fff1e0;}}
tr.pending td:first-child{{box-shadow:inset 4px 0 0 var(--pending);}}
.banner{{border-radius:var(--radius);padding:11px 14px;margin-bottom:14px;font-size:13px;font-weight:700;}}
.banner.bad{{background:#fdecea;border:1px solid var(--bad);color:#7a231c;}}
.banner.pending{{background:#fff1e0;border:1px solid var(--pending);color:#8a3d00;}}
.empty{{color:var(--muted);}}
@media(max-width:760px){{.wrap{{padding:12px;}}.header{{display:block;}}.meta{{text-align:left;margin-top:10px;}}.actions{{justify-content:flex-start;}}.doc{{padding:14px;}}}}
</style>
</head>
<body>
<div class="wrap">
  <header class="header">
    <div>
      <div class="brand">Vista Word · Post-ventas</div>
      <h1>{escape(docx_path.name)}</h1>
    </div>
    <div class="meta">
      Generado: {escape(generated)}<br>
      <div class="actions">
        <a class="btn" href="{escape(index_href, quote=True)}">Panel</a>
        <a class="btn" href="{escape(original_href, quote=True)}">Abrir Word</a>
      </div>
    </div>
  </header>
  {error}
  <main class="doc">{content}</main>
</div>
</body>
</html>
"""
    preview_path.write_text(html, encoding="utf-8")
    return preview_path, total_pending, total_rows


def attach_word_previews(obras: list[dict]) -> None:
    for obra in obras:
        previews = {}
        pending_by_word = {}
        for word in obra["words"]:
            preview, pending_count, rows_considered = write_preview(word)
            if preview:
                previews[word] = preview
                pending_by_word[word] = pending_count
        obra["previews"] = previews
        obra["pending_count"] = pending_by_word.get(obra["main_word"], 0) if obra.get("main_word") else 0
        obra["vencido"] = is_vencido(obra)


def render_latest_files(paths: list[Path], previews: dict[Path, Path]) -> str:
    if not paths:
        return '<li><span>Sin archivos en la carpeta</span></li>'
    rows = []
    for path in paths:
        href_path = previews.get(path, path)
        suffix = "Vista HTML" if path in previews else fmt_dt(path.stat().st_mtime)
        rows.append(
            '<li>'
            f'<a href="{escape(url_for(href_path), quote=True)}">{escape(path.name)}</a>'
            f'<span>{escape(suffix)} · {escape(fmt_dt(path.stat().st_mtime))}</span>'
            '</li>'
        )
    return "".join(rows)


def render_card(obra: dict) -> str:
    last_file = obra["last_file"]
    last_name = last_file.name if last_file else "Sin archivos"
    recent = is_recent(obra)
    word_preview = obra["previews"].get(obra["main_word"]) if obra["main_word"] else None
    search_text = " ".join([
        obra["name"],
        obra["raw_name"],
        last_name,
        " ".join(p.name for p in obra["latest_files"]),
    ]).lower()

    pending_count = obra.get("pending_count", 0)
    vencido = obra.get("vencido", False)

    badges = []
    if pending_count:
        plural = "s" if pending_count != 1 else ""
        badges.append(f'<span class="badge pending">{pending_count} pendiente{plural}</span>')
    if vencido:
        badges.append('<span class="badge vencido" title="Han pasado más de 2 años desde la última incidencia real: la postventa de esta obra ha vencido. Lo que se registre aquí a partir de ahora son casos aislados / gestión administrativa, ya no postventa.">VENCIDO</span>')
    if recent:
        badges.append('<span class="badge ok">Reciente</span>')
    if obra["pdfs"]:
        badges.append('<span class="badge">PDF</span>')
    if obra["words"]:
        badges.append('<span class="badge">Word</span>')
    if obra["images"]:
        badges.append('<span class="badge">Fotos</span>')

    return f"""
<article class="obra" data-search="{escape(search_text, quote=True)}"
  data-recent="{str(recent).lower()}" data-vencido="{str(vencido).lower()}" data-pdf="{str(bool(obra['pdfs'])).lower()}"
  data-word="{str(bool(obra['words'])).lower()}" data-images="{str(bool(obra['images'])).lower()}">
  <div class="card-top">
    <div>
      <div class="eyebrow">Contrato postventa</div>
      <h2>{escape(obra["name"])}</h2>
    </div>
    <div class="age">{escape(days_text(obra["last_ts"]))}</div>
  </div>
  <div class="badges">{"".join(badges) or '<span class="badge muted">Sin clasificar</span>'}</div>
  <div class="last">
    <span>Último archivo</span>
    <strong>{escape(fmt_dt(obra["last_ts"]))}</strong>
    <em>{escape(last_name)}</em>
  </div>
  <div class="metrics">
    <div><b>{len(obra["files"])}</b><span>Archivos</span></div>
    <div><b>{len(obra["pdfs"])}</b><span>PDF</span></div>
    <div><b>{len(obra["words"])}</b><span>Word</span></div>
    <div><b>{len(obra["images"])}</b><span>Fotos</span></div>
  </div>
  <div class="actions">
    <a class="btn primary" href="{escape(url_for(obra["folder"]), quote=True)}/">Abrir carpeta</a>
    {action_link("Ver Word", word_preview)}
    {action_link("Abrir Word", obra["main_word"])}
    {action_link("Último PDF", obra["last_pdf"])}
  </div>
  <details>
    <summary>Últimos documentos</summary>
    <ul class="docs">{render_latest_files(obra["latest_files"], obra["previews"])}</ul>
  </details>
</article>"""


def build_html(obras: list[dict]) -> str:
    total_pdfs = sum(len(o["pdfs"]) for o in obras)
    total_words = sum(len(o["words"]) for o in obras)
    total_images = sum(len(o["images"]) for o in obras)
    total_pending = sum(o.get("pending_count", 0) for o in obras)
    total_vencidas = sum(1 for o in obras if o.get("vencido"))
    latest_ts = max((o["last_ts"] for o in obras), default=None)
    cards = "\n".join(render_card(o) for o in obras)

    return f"""<!DOCTYPE html>
<html lang="es">
<head>
<meta charset="UTF-8">
<meta name="viewport" content="width=device-width, initial-scale=1.0">
<title>Post-ventas Sagarde</title>
<style>
:root{{--bg:#f4f6f9;--card:#fff;--header:#0b1f3a;--header2:#123a63;--text:#1c2733;--muted:#647184;--accent:#f5a524;--ok:#2e9e5b;--warn:#e07b1a;--bad:#d9483c;--pending:#e2680a;--line:#e3e7ee;--radius:10px;}}
*{{margin:0;padding:0;box-sizing:border-box;}}
body{{font-family:-apple-system,BlinkMacSystemFont,'Segoe UI',Roboto,sans-serif;background:var(--bg);color:var(--text);line-height:1.45;}}
.wrap{{max-width:1320px;margin:0 auto;padding:22px;}}
.top{{background:linear-gradient(120deg,var(--header),var(--header2));color:#fff;border-bottom:4px solid #b42318;}}
.top-inner{{max-width:1320px;margin:auto;padding:18px 28px;display:flex;align-items:center;gap:20px;}}
.logo{{width:min(290px,22vw);height:auto;object-fit:contain;border-radius:9px;box-shadow:0 0 0 3px #b42318,0 6px 28px rgba(0,0,0,.5);}}
.identity{{min-width:0;}}.identity strong{{font-size:21px;display:block;}}.identity span{{font-size:12px;color:#c7d3e3;}}
.top-actions{{margin-left:auto;display:flex;gap:8px;}}
.top-actions a{{text-decoration:none;border:1px solid rgba(255,255,255,.35);color:#fff;padding:9px 12px;border-radius:5px;font-size:13px;font-weight:600;}}
.intro{{display:flex;justify-content:space-between;align-items:center;gap:20px;flex-wrap:wrap;margin-bottom:14px;}}
.search-wrap{{width:min(360px,100%);position:relative;}}
.search-symbol{{position:absolute;right:14px;top:11px;color:#667085;}}
.sub{{font-size:13px;color:var(--muted);margin:0;}}
.stats{{display:grid;grid-template-columns:repeat(auto-fit,minmax(130px,1fr));gap:10px;margin-bottom:14px;}}
.stat{{background:var(--card);border-radius:var(--radius);padding:13px 16px;box-shadow:0 1px 3px rgba(0,0,0,.07);}}
.stat-value{{font-size:23px;font-weight:800;}}
.stat-value.pending{{color:var(--pending);}}
.stat-label{{font-size:11px;color:var(--muted);text-transform:uppercase;letter-spacing:.5px;margin-top:2px;}}
.toolbar{{display:flex;gap:10px;align-items:center;justify-content:space-between;flex-wrap:wrap;margin-bottom:14px;}}
.search{{width:100%;height:42px;padding:0 42px 0 14px;border:1px solid #98a2b3;border-radius:6px;background:#fff;font:inherit;color:var(--text);}}
.filters{{display:flex;gap:7px;flex-wrap:wrap;}}
.filters button{{border:1px solid var(--line);background:#fff;color:var(--text);border-radius:8px;padding:9px 12px;font-size:13px;font-weight:650;cursor:pointer;}}
.filters button.active{{background:var(--header);border-color:var(--header);color:#fff;}}
.count{{font-size:12px;color:var(--muted);font-weight:650;}}
.grid{{display:grid;grid-template-columns:repeat(auto-fill,minmax(315px,1fr));gap:14px;align-items:start;}}
.obra{{background:var(--card);border-radius:var(--radius);padding:18px 20px;box-shadow:0 1px 3px rgba(0,0,0,.08);border-top:4px solid var(--accent);}}
.card-top{{display:flex;justify-content:space-between;gap:12px;align-items:flex-start;}}
.eyebrow{{font-size:10.5px;color:var(--muted);text-transform:uppercase;letter-spacing:.5px;font-weight:800;margin-bottom:4px;}}
.obra h2{{font-size:16px;line-height:1.25;font-weight:760;}}
.age{{font-size:11.5px;color:var(--muted);white-space:nowrap;margin-top:1px;}}
.badges{{display:flex;gap:6px;flex-wrap:wrap;margin:12px 0;}}
.badge{{font-size:11px;font-weight:750;border-radius:20px;background:#eef2f7;color:var(--header2);padding:3px 8px;}}
.badge.ok{{background:#e8f6ee;color:var(--ok);}}
.badge.muted{{color:var(--muted);}}
.badge.pending{{background:#fff1e0;color:var(--pending);}}
.badge.vencido{{background:#fdecea;color:var(--bad);cursor:help;}}
.last{{border-top:1px solid var(--line);border-bottom:1px solid var(--line);padding:11px 0;margin-bottom:12px;}}
.last span{{display:block;font-size:11px;text-transform:uppercase;letter-spacing:.5px;color:var(--muted);font-weight:800;}}
.last strong{{display:block;font-size:18px;margin-top:2px;}}
.last em{{display:block;font-style:normal;color:var(--muted);font-size:12px;margin-top:3px;overflow:hidden;text-overflow:ellipsis;white-space:nowrap;}}
.metrics{{display:grid;grid-template-columns:repeat(4,1fr);gap:8px;margin-bottom:12px;}}
.metrics div{{background:#f8f9fb;border-radius:8px;padding:9px 8px;text-align:center;}}
.metrics b{{display:block;font-size:18px;}}
.metrics span{{font-size:10.5px;color:var(--muted);text-transform:uppercase;letter-spacing:.4px;}}
.actions{{display:flex;gap:7px;flex-wrap:wrap;margin-bottom:10px;}}
.btn{{display:inline-flex;align-items:center;justify-content:center;text-decoration:none;border:1px solid var(--line);background:#fff;color:var(--header2);font-weight:750;font-size:12.5px;border-radius:7px;padding:7px 10px;min-height:32px;}}
.btn.primary{{background:var(--header);border-color:var(--header);color:#fff;}}
.btn.disabled{{color:#a2adbb;background:#f4f6f9;}}
details{{border-top:1px solid var(--line);padding-top:9px;}}
summary{{cursor:pointer;font-size:12.5px;font-weight:750;color:var(--header2);}}
.docs{{list-style:none;margin-top:8px;display:grid;gap:7px;}}
.docs li{{display:grid;gap:2px;font-size:12px;min-width:0;}}
.docs a{{color:var(--header2);text-decoration:none;overflow:hidden;text-overflow:ellipsis;white-space:nowrap;}}
.docs a:hover{{text-decoration:underline;}}
.docs span{{color:var(--muted);font-size:11px;}}
.empty{{display:none;background:#fff;border-radius:var(--radius);padding:30px;text-align:center;color:var(--muted);}}
.footer{{text-align:center;font-size:12px;color:var(--muted);padding:18px 0;}}
@media(max-width:900px){{.logo{{width:190px;}}}}
@media(max-width:760px){{.wrap{{padding:14px;}}.top-actions{{display:none;}}.logo{{width:150px;}}.metrics{{grid-template-columns:repeat(2,1fr);}}}}
</style>
</head>
<body>
<header class="top"><div class="top-inner">
<img class="logo" src="logo_sagarde.jpg" alt="Sagarde">
<div class="identity"><strong>Post-ventas</strong><span>Incidencias, partes resueltos y matrices por obra</span></div>
<nav class="top-actions">
  <a href="../index.html">&#8962; Portal</a>
  <a href="../SAGARDE%20OBRAS%20ABIERTAS/index.html">Obras abiertas</a>
  <a href="./index.html">Post-ventas</a>
  <a href="../MANTENIMIENTOS/index.html">Mantenimientos</a>
  <a href="../SAGARDE%20(OLD)/OBRAS%20CERRADAS/index.html">Obras cerradas</a>
</nav>
</div></header>
<div class="wrap">
  <div class="intro">
    <p class="sub">{len(obras)} contrato(s) de postventa · incidencias activas primero, luego por último archivo · generado {datetime.now().strftime("%d/%m/%Y %H:%M")}</p>
    <label class="search-wrap"><input id="search" class="search" type="search" placeholder="Buscar obra, archivo, código o incidencia"><span class="search-symbol">&#9906;</span></label>
  </div>

  <section class="stats">
    {stat_block("Contratos", len(obras))}
    {stat_block("Pendientes", total_pending, "pending" if total_pending else "")}
    {stat_block("Vencidas", total_vencidas)}
    {stat_block("PDF", total_pdfs)}
    {stat_block("Word", total_words)}
    {stat_block("Fotos", total_images)}
    {stat_block("Última actividad", fmt_dt(latest_ts))}
  </section>

  <div class="toolbar">
    <div class="filters">
      <button class="active" data-filter="all">Todas</button>
      <button data-filter="recent">Recientes</button>
      <button data-filter="vencido">Vencidas</button>
      <button data-filter="pdf">Con PDF</button>
      <button data-filter="word">Con Word</button>
      <button data-filter="images">Con fotos</button>
    </div>
    <div id="count" class="count"></div>
  </div>

  <main id="grid" class="grid">{cards}</main>
  <div id="empty" class="empty">No hay coincidencias con el filtro actual.</div>
  <div class="footer">Para actualizar este panel: doble clic en <b>Actualizar_Postventas.bat</b>.</div>
</div>
<script>
const search = document.getElementById('search');
const buttons = [...document.querySelectorAll('.filters button')];
const cards = [...document.querySelectorAll('.obra')];
const count = document.getElementById('count');
const empty = document.getElementById('empty');
let currentFilter = 'all';

function matchesFilter(card) {{
  if (currentFilter === 'all') return true;
  return card.dataset[currentFilter] === 'true';
}}

function applyFilters() {{
  const q = search.value.trim().toLowerCase();
  let visible = 0;
  cards.forEach(card => {{
    const ok = matchesFilter(card) && (!q || card.dataset.search.includes(q));
    card.style.display = ok ? '' : 'none';
    if (ok) visible += 1;
  }});
  count.textContent = `${{visible}} de ${{cards.length}}`;
  empty.style.display = visible ? 'none' : 'block';
}}

buttons.forEach(btn => {{
  btn.addEventListener('click', () => {{
    buttons.forEach(b => b.classList.remove('active'));
    btn.classList.add('active');
    currentFilter = btn.dataset.filter;
    applyFilters();
  }});
}});
search.addEventListener('input', applyFilters);
applyFilters();
</script>
</body>
</html>
"""


def escribir_resumen_json(obras: list[dict]) -> None:
    """
    Publica un resumen en JSON para que lo lea el portal raiz
    (COPIA SEGURIDAD SAGARDE/sagarde_portal.py) sin reimplementar el
    escaneo de carpetas. Contrato de solo lectura, igual que
    resumen_obras.json del sistema de Obras Abiertas.
    """
    ahora = datetime.now().timestamp()
    resumen = {
        "generado": datetime.now().strftime("%d/%m/%Y %H:%M"),
        "generado_ts": ahora,
        "totales": {
            "n_contratos": len(obras),
            "n_recientes": sum(1 for o in obras if is_recent(o)),
            "n_pdfs": sum(len(o["pdfs"]) for o in obras),
            "n_words": sum(len(o["words"]) for o in obras),
        },
        "contratos": [
            {
                "nombre": o["name"],
                "carpeta": o["raw_name"],
                "n_archivos": len(o["files"]),
                "reciente": is_recent(o),
                "vencido": o.get("vencido", False),
                "ultimo_archivo_ts": o["last_ts"],
                "ultima_incidencia_ts": o.get("last_incident_ts"),
                "href": "POST-VENTAS/" + url_for(o["folder"]) + "/",
            }
            for o in obras
        ],
    }
    RESUMEN_JSON.write_text(json.dumps(resumen, ensure_ascii=False, indent=2), encoding="utf-8")
    print(f"Resumen JSON: {RESUMEN_JSON}")


def reorder_by_pending(obras: list[dict]) -> list[dict]:
    """Ordena el panel en tres bloques:
    1) obras con incidencias activas (pending_count > 0) primero, de más a
       menos pendientes (a igualdad, por archivo más reciente) -- son las
       que requieren atención ahora mismo;
    2) obras sin pendientes pero con la postventa todavia en plazo (no
       vencida), en el orden por archivo más reciente que ya trae
       scan_obras();
    3) obras vencidas (más de 2 años sin actividad real) al final, porque
       cualquier cosa que se registre ahi ya no es postventa propiamente
       dicha."""
    con_pendientes = [o for o in obras if o.get("pending_count", 0) > 0]
    resto = [o for o in obras if not o.get("pending_count", 0)]
    con_pendientes.sort(key=lambda o: (-o.get("pending_count", 0), -o["last_ts"]))
    en_plazo = [o for o in resto if not o.get("vencido")]
    vencidas = [o for o in resto if o.get("vencido")]
    return con_pendientes + en_plazo + vencidas


def main() -> None:
    obras = scan_obras()
    attach_word_previews(obras)
    obras = reorder_by_pending(obras)
    INDEX_PATH.write_text(build_html(obras), encoding="utf-8")
    escribir_resumen_json(obras)
    print(f"Index generado: {INDEX_PATH}")
    print(f"Contratos: {len(obras)}")


if __name__ == "__main__":
    main()
